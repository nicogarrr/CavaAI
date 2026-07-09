import hashlib
import os
import re
import tempfile
from dataclasses import dataclass, field
from datetime import UTC, datetime
from html.parser import HTMLParser
from io import BytesIO
from pathlib import Path
from urllib.parse import urlparse

import httpx
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.models import Company, Document, DocumentChunk
from app.services.document_store import DocumentStore


MAX_DOCUMENT_BYTES = 15 * 1024 * 1024
SUPPORTED_EXTENSIONS = {".txt", ".md", ".html", ".htm", ".pdf", ".docx", ".xlsx", ".csv", ".tsv"}


@dataclass
class ParsedBlock:
    text: str
    metadata: dict = field(default_factory=dict)


@dataclass
class ParsedDocument:
    blocks: list[ParsedBlock]
    parser: str
    warnings: list[str] = field(default_factory=list)


class _HTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []

    def handle_data(self, data: str) -> None:
        stripped = " ".join(data.split())
        if stripped:
            self._parts.append(stripped)

    def text(self) -> str:
        return "\n".join(self._parts)


def _compact(text: str) -> str:
    return re.sub(r"\n{3,}", "\n\n", text.replace("\r\n", "\n").replace("\r", "\n")).strip()


def _extension(filename: str, content_type: str | None = None) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix:
        return suffix
    content_type = (content_type or "").lower()
    if "pdf" in content_type:
        return ".pdf"
    if "spreadsheet" in content_type or "excel" in content_type:
        return ".xlsx"
    if "word" in content_type:
        return ".docx"
    if "html" in content_type:
        return ".html"
    return ".txt"


class DocumentIngestionService:
    def ingest_bytes(
        self,
        db: Session,
        *,
        ticker: str,
        title: str,
        content: bytes,
        filename: str,
        source_type: str,
        source_url: str | None = None,
        content_type: str | None = None,
        published_at: datetime | None = None,
    ) -> dict:
        if not content:
            raise ValueError("Document is empty")
        if len(content) > MAX_DOCUMENT_BYTES:
            raise ValueError("Document exceeds 15MB local ingestion limit")

        company = db.scalar(select(Company).where(Company.ticker == ticker.upper()))
        if not company:
            raise ValueError(f"Company {ticker.upper()} not found")

        checksum = hashlib.sha256(content).hexdigest()
        duplicate = db.scalar(
            select(Document).where(Document.company_id == company.id, Document.checksum == checksum)
        )
        if duplicate:
            chunk_count = db.scalar(
                select(DocumentChunk.id).where(DocumentChunk.document_id == duplicate.id).limit(1)
            )
            return {
                "status": "duplicate",
                "ticker": company.ticker,
                "document_id": duplicate.id,
                "chunks": 0 if chunk_count is None else len(duplicate.chunks),
                "checksum": checksum,
                "parser": duplicate.metadata_.get("parser", "unknown"),
                "storage_uri": duplicate.storage_uri,
                "warnings": ["Document with same checksum already exists for this company."],
            }

        ext = _extension(filename, content_type)
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported document extension: {ext}")

        parsed = self._parse(content, filename, ext, content_type)
        text = "\n\n".join(block.text for block in parsed.blocks if block.text.strip())
        if len(text.strip()) < 20:
            raise ValueError("Document parser produced too little text")

        storage_uri = DocumentStore().put_bytes_local(
            company.ticker,
            source_type or "manual",
            f"{checksum[:12]}-{filename}",
            content,
            tenant_id=db.info.get("tenant_id"),
        )

        document = Document(
            company_id=company.id,
            title=title,
            source_type=source_type,
            source_url=source_url,
            storage_uri=storage_uri,
            published_at=published_at or datetime.now(UTC),
            checksum=checksum,
            metadata_={
                "parser": parsed.parser,
                "filename": filename,
                "content_type": content_type,
                "extension": ext,
                "raw_size_bytes": len(content),
                "block_count": len(parsed.blocks),
                "warnings": parsed.warnings,
                "docling_opt_in": os.getenv("CAVAAI_USE_DOCLING") == "1",
            },
        )
        db.add(document)
        db.flush()

        chunks = self._chunk_blocks(parsed.blocks, checksum, parsed.parser, filename, source_url)
        for index, chunk in enumerate(chunks):
            db.add(
                DocumentChunk(
                    document_id=document.id,
                    chunk_index=index,
                    text=chunk["text"],
                    token_count=len(chunk["text"].split()),
                    metadata_=chunk["metadata"],
                )
            )

        db.commit()
        db.refresh(document)

        if os.getenv("CAVAAI_ENABLE_VECTOR_INGEST") == "1":
            try:
                from app.services.rag import RAGIndex

                rag_result = RAGIndex().ingest_document(db, document)
            except Exception as exc:
                rag_result = {"chunks_indexed": 0, "error": str(exc)}
        else:
            rag_result = {"chunks_indexed": 0, "skipped": "Set CAVAAI_ENABLE_VECTOR_INGEST=1 to index Qdrant."}

        if os.getenv("CAVAAI_ENABLE_AUTO_RESEARCH", "1") == "1":
            try:
                from app.services.claim_intelligence_service import (
                    ClaimIntelligenceService,
                )

                intelligence_result = ClaimIntelligenceService().scan_document(
                    db, document, auto_apply=True
                )
            except Exception as exc:
                intelligence_result = {
                    "status": "failed",
                    "error": str(exc),
                    "document_id": document.id,
                }
        else:
            intelligence_result = {
                "status": "disabled",
                "document_id": document.id,
            }

        return {
            "status": "ingested",
            "ticker": company.ticker,
            "document_id": document.id,
            "chunks": len(chunks),
            "checksum": checksum,
            "parser": parsed.parser,
            "storage_uri": storage_uri,
            "warnings": parsed.warnings,
            "rag": rag_result,
            "intelligence": intelligence_result,
        }

    def ingest_url(
        self,
        db: Session,
        *,
        ticker: str,
        title: str,
        url: str,
        source_type: str,
    ) -> dict:
        parsed_url = urlparse(url)
        if parsed_url.scheme not in {"http", "https"}:
            raise ValueError("Only http(s) URLs can be ingested")

        with httpx.Client(timeout=20, follow_redirects=True) as client:
            response = client.get(url)
            response.raise_for_status()
            content = response.content
            content_type = response.headers.get("content-type")

        filename = Path(parsed_url.path).name or f"{parsed_url.netloc}.html"
        return self.ingest_bytes(
            db,
            ticker=ticker,
            title=title,
            content=content,
            filename=filename,
            source_type=source_type,
            source_url=url,
            content_type=content_type,
        )

    def _parse(self, content: bytes, filename: str, ext: str, content_type: str | None) -> ParsedDocument:
        docling = self._parse_with_docling(content, filename, ext)
        if docling and docling.blocks:
            return docling
        docling_warnings = docling.warnings if docling else []
        parsed = self._parse_native(content, ext)
        parsed.warnings.extend(docling_warnings)
        return parsed

    def _parse_native(self, content: bytes, ext: str) -> ParsedDocument:
        if ext in {".txt", ".md", ".csv", ".tsv"}:
            return self._parse_text(content, ext)
        if ext in {".html", ".htm"}:
            return self._parse_html(content)
        if ext == ".pdf":
            return self._parse_pdf(content)
        if ext == ".docx":
            return self._parse_docx(content)
        if ext == ".xlsx":
            return self._parse_xlsx(content)
        return self._parse_text(content, ext)

    def _parse_with_docling(self, content: bytes, filename: str, ext: str) -> ParsedDocument | None:
        if os.getenv("CAVAAI_USE_DOCLING") != "1":
            return None
        try:
            from docling.document_converter import DocumentConverter
        except Exception:
            return ParsedDocument(
                blocks=[],
                parser="docling_unavailable",
                warnings=["CAVAAI_USE_DOCLING=1 but docling is not installed; native parser fallback used."],
            )

        suffix = ext if ext.startswith(".") else f".{ext}"
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as handle:
            handle.write(content)
            temp_path = handle.name
        try:
            result = DocumentConverter().convert(temp_path)
            markdown = result.document.export_to_markdown()
            if markdown.strip():
                return ParsedDocument(
                    blocks=[ParsedBlock(text=_compact(markdown), metadata={"docling": True})],
                    parser="docling",
                )
        except Exception as exc:
            return ParsedDocument(
                blocks=[],
                parser="docling_failed",
                warnings=[f"Docling parser failed: {exc}. Native parser fallback used."],
            )
        finally:
            try:
                os.unlink(temp_path)
            except OSError:
                pass
        return None

    def _parse_text(self, content: bytes, ext: str) -> ParsedDocument:
        text = content.decode("utf-8", errors="replace")
        return ParsedDocument(
            blocks=[ParsedBlock(text=_compact(text), metadata={"format": ext.lstrip(".")})],
            parser="native_text",
        )

    def _parse_html(self, content: bytes) -> ParsedDocument:
        parser = _HTMLTextExtractor()
        parser.feed(content.decode("utf-8", errors="replace"))
        return ParsedDocument(
            blocks=[ParsedBlock(text=_compact(parser.text()), metadata={"format": "html"})],
            parser="native_html",
        )

    def _parse_pdf(self, content: bytes) -> ParsedDocument:
        from PyPDF2 import PdfReader

        reader = PdfReader(BytesIO(content))
        blocks = []
        for page_index, page in enumerate(reader.pages):
            text = _compact(page.extract_text() or "")
            if text:
                blocks.append(ParsedBlock(text=text, metadata={"page": page_index + 1}))
        return ParsedDocument(blocks=blocks, parser="pypdf2")

    def _parse_docx(self, content: bytes) -> ParsedDocument:
        from docx import Document as DocxDocument

        document = DocxDocument(BytesIO(content))
        blocks = []
        for index, paragraph in enumerate(document.paragraphs):
            text = _compact(paragraph.text)
            if text:
                blocks.append(ParsedBlock(text=text, metadata={"paragraph": index + 1}))
        return ParsedDocument(blocks=blocks, parser="python_docx")

    def _parse_xlsx(self, content: bytes) -> ParsedDocument:
        from openpyxl import load_workbook

        workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
        blocks = []
        for sheet in workbook.worksheets:
            rows = []
            for row in sheet.iter_rows(values_only=True):
                values = [str(value) for value in row if value is not None]
                if values:
                    rows.append(" | ".join(values))
            if rows:
                blocks.append(ParsedBlock(text="\n".join(rows), metadata={"sheet": sheet.title}))
        return ParsedDocument(blocks=blocks, parser="openpyxl")

    def _chunk_blocks(
        self,
        blocks: list[ParsedBlock],
        checksum: str,
        parser: str,
        filename: str,
        source_url: str | None,
        max_chars: int = 2500,
    ) -> list[dict]:
        chunks: list[dict] = []
        current_text: list[str] = []
        current_meta: list[dict] = []

        def flush() -> None:
            if not current_text:
                return
            text = _compact("\n\n".join(current_text))
            if not text:
                current_text.clear()
                current_meta.clear()
                return
            chunks.append(
                {
                    "text": text,
                    "metadata": {
                        "checksum": checksum,
                        "parser": parser,
                        "filename": filename,
                        "source_url": source_url,
                        "block_metadata": list(current_meta),
                        "chunk_sha256": hashlib.sha256(text.encode("utf-8")).hexdigest(),
                    },
                }
            )
            current_text.clear()
            current_meta.clear()

        for block in blocks:
            text = _compact(block.text)
            if not text:
                continue
            if len(text) > max_chars:
                flush()
                words = text.split()
                piece: list[str] = []
                for word in words:
                    if sum(len(item) + 1 for item in piece) + len(word) > max_chars:
                        current_text.append(" ".join(piece))
                        current_meta.append(block.metadata)
                        flush()
                        piece = []
                    piece.append(word)
                if piece:
                    current_text.append(" ".join(piece))
                    current_meta.append(block.metadata)
                continue
            if sum(len(item) + 2 for item in current_text) + len(text) > max_chars:
                flush()
            current_text.append(text)
            current_meta.append(block.metadata)
        flush()
        return chunks
