from fastapi import FastAPI, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from typing import Optional
from modules.fmp import (
    fetch_income_statement,
    fetch_balance_sheet,
    fetch_cash_flow,
    fetch_financial_growth,
    fetch_ratios_ttm,
    fetch_dcf,
    fetch_enterprise_value,
    fetch_key_metrics_ttm,
    fetch_financial_scores,
    fetch_owner_earnings,
    fetch_price_target_consensus,
    fetch_grades_consensus,
    fetch_stock_peers,
    # Priority APIs
    fetch_earnings_transcript,
    fetch_insider_trading,
    fetch_treasury_rates,
    fetch_analyst_estimates,
    fetch_treasury_rates,
    fetch_analyst_estimates,
    fetch_press_releases,
    # Market Movers & Screener
    fetch_biggest_gainers,
    fetch_biggest_losers,
    fetch_most_actives,
    fetch_stock_screener,
    fetch_earnings_transcripts_list,
    fetch_earnings_transcript,
    fetch_fmp_articles,
    fetch_general_news
)

app = FastAPI(title="FMP Data Engine", version="1.0.0")

# CORS for Next.js frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
async def root():
    return {"status": "ok", "service": "FMP Data Engine"}

@app.get("/fundamentals/{symbol}")
async def get_fundamentals(symbol: str, period: str = "annual"):
    """Get combined financial statements (Income, Balance, CashFlow)"""
    symbol = symbol.upper()
    try:
        income = fetch_income_statement(symbol, period)
        balance = fetch_balance_sheet(symbol, period)
        cashflow = fetch_cash_flow(symbol, period)
        
        if any('error' in d for d in [income, balance, cashflow] if isinstance(d, dict)):
            raise HTTPException(status_code=500, detail="Error fetching financial statements")
        
        return {
            "symbol": symbol,
            "period": period,
            "income": income,
            "balance": balance,
            "cashflow": cashflow
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/financial-growth/{symbol}")
async def get_financial_growth(symbol: str):
    """Get pre-calculated growth metrics (revenue, EPS, FCF growth)"""
    symbol = symbol.upper()
    try:
        data = fetch_financial_growth(symbol)
        if isinstance(data, dict) and 'error' in data:
            raise HTTPException(status_code=500, detail=data['error'])
        return {"symbol": symbol, "growth": data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/ratios-ttm/{symbol}")
async def get_ratios_ttm(symbol: str):
    """Get TTM ratios (PER, ROIC, ROE, P/S, P/B, etc.)"""
    symbol = symbol.upper()
    try:
        data = fetch_ratios_ttm(symbol)
        if isinstance(data, dict) and 'error' in data:
            raise HTTPException(status_code=500, detail=data['error'])
        return {"symbol": symbol, "ratios": data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/dcf/{symbol}")
async def get_dcf(symbol: str):
    """Get Discounted Cash Flow valuation (intrinsic value)"""
    symbol = symbol.upper()
    try:
        data = fetch_dcf(symbol)
        if isinstance(data, dict) and 'error' in data:
            raise HTTPException(status_code=500, detail=data['error'])
        return {"symbol": symbol, "dcf": data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/enterprise-value/{symbol}")
async def get_enterprise_value(symbol: str):
    """Get Enterprise Value data (for EV/EBITDA, EV/FCF)"""
    symbol = symbol.upper()
    try:
        data = fetch_enterprise_value(symbol)
        if isinstance(data, dict) and 'error' in data:
            raise HTTPException(status_code=500, detail=data['error'])
        return {"symbol": symbol, "enterpriseValue": data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/key-metrics-ttm/{symbol}")
async def get_key_metrics_ttm(symbol: str):
    """Get Key Metrics TTM (ROE, ROIC, EV/EBITDA, Graham Number, etc.)"""
    symbol = symbol.upper()
    try:
        data = fetch_key_metrics_ttm(symbol)
        if isinstance(data, dict) and 'error' in data:
            raise HTTPException(status_code=500, detail=data['error'])
        return {"symbol": symbol, "keyMetrics": data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/financial-scores/{symbol}")
async def get_financial_scores_endpoint(symbol: str):
    """Get Financial Scores (Altman Z-Score + Piotroski Score)"""
    symbol = symbol.upper()
    try:
        data = fetch_financial_scores(symbol)
        if isinstance(data, dict) and 'error' in data:
            raise HTTPException(status_code=500, detail=data['error'])
        # Return data directly (array) as frontend expects
        return data
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/owner-earnings/{symbol}")
async def get_owner_earnings(symbol: str):
    """Get Owner Earnings (Buffett's preferred metric)"""
    symbol = symbol.upper()
    try:
        data = fetch_owner_earnings(symbol)
        if isinstance(data, dict) and 'error' in data:
            raise HTTPException(status_code=500, detail=data['error'])
        return {"symbol": symbol, "ownerEarnings": data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/price-target/{symbol}")
async def get_price_target(symbol: str):
    """Get Analyst Price Target Consensus"""
    symbol = symbol.upper()
    try:
        data = fetch_price_target_consensus(symbol)
        if isinstance(data, dict) and 'error' in data:
            raise HTTPException(status_code=500, detail=data['error'])
        return {"symbol": symbol, "priceTarget": data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/grades/{symbol}")
async def get_grades(symbol: str):
    """Get Stock Grades Consensus (Buy/Hold/Sell)"""
    symbol = symbol.upper()
    try:
        data = fetch_grades_consensus(symbol)
        if isinstance(data, dict) and 'error' in data:
            raise HTTPException(status_code=500, detail=data['error'])
        return {"symbol": symbol, "grades": data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/peers/{symbol}")
async def get_peers(symbol: str):
    """Get Stock Peers for comparison"""
    symbol = symbol.upper()
    try:
        data = fetch_stock_peers(symbol)
        if isinstance(data, dict) and 'error' in data:
            raise HTTPException(status_code=500, detail=data['error'])
        return {"symbol": symbol, "peers": data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ============================================================================
# PRIORITY APIs - AI/RAG, Trading Signals, WACC
# ============================================================================

@app.get("/earnings-transcript/{symbol}")
async def get_earnings_transcript(
    symbol: str,
    year: Optional[int] = Query(None, description="Year of earnings call"),
    quarter: Optional[int] = Query(None, description="Quarter (1-4)")
):
    """Get Earnings Call Transcript for AI/RAG analysis"""
    symbol = symbol.upper()
    try:
        data = fetch_earnings_transcript(symbol, year, quarter)
        if isinstance(data, dict) and 'error' in data:
            raise HTTPException(status_code=500, detail=data['error'])
        return {"symbol": symbol, "transcripts": data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/insider-trading/{symbol}")
async def get_insider_trading(
    symbol: str,
    limit: int = Query(50, description="Number of transactions to return")
):
    """Get Insider Trading data (CEO/CFO buy/sell signals)"""
    symbol = symbol.upper()
    try:
        data = fetch_insider_trading(symbol, limit)
        if isinstance(data, dict) and 'error' in data:
            raise HTTPException(status_code=500, detail=data['error'])
        return {"symbol": symbol, "insiderTrades": data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/treasury-rates")
async def get_treasury_rates():
    """Get Treasury Rates (10Y for Risk-Free Rate in WACC)"""
    try:
        data = fetch_treasury_rates()
        if isinstance(data, dict) and 'error' in data:
            raise HTTPException(status_code=500, detail=data['error'])
        return {"treasuryRates": data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/analyst-estimates/{symbol}")
async def get_analyst_estimates(
    symbol: str,
    period: str = Query("annual", description="annual or quarter"),
    limit: int = Query(5, description="Number of estimates to return")
):
    """Get Analyst Estimates (Future EPS/Revenue projections)"""
    symbol = symbol.upper()
    try:
        data = fetch_analyst_estimates(symbol, period, limit)
        if isinstance(data, dict) and 'error' in data:
            raise HTTPException(status_code=500, detail=data['error'])
        return {"symbol": symbol, "estimates": data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/press-releases/{symbol}")
async def get_press_releases(
    symbol: str,
    limit: int = Query(20, description="Number of releases to return")
):
    """Get Official Press Releases"""
    symbol = symbol.upper()
    try:
        data = fetch_press_releases(symbol, limit)
        if isinstance(data, dict) and 'error' in data:
            print(f"Error fetching press releases for {symbol}: {data['error']}")
            return {"symbol": symbol, "pressReleases": []}
        
        # Ensure data is a list. If it's a dict (like an error message from FMP without 'error' key), treat as empty
        if not isinstance(data, list):
            print(f"Unexpected response format for press releases {symbol}: {data}")
            return {"symbol": symbol, "pressReleases": []}

        return {"symbol": symbol, "pressReleases": data}
    except Exception as e:
        print(f"Exception in get_press_releases: {str(e)}")
        return {"symbol": symbol, "pressReleases": []}

@app.get("/health")
async def health_check():
    return {"status": "healthy"}

@app.get("/market-movers/gainers")
async def get_biggest_gainers():
    """Get Biggest Stock Gainers"""
    try:
        data = fetch_biggest_gainers()
        if isinstance(data, dict) and 'error' in data:
            raise HTTPException(status_code=500, detail=data['error'])
        return {"gainers": data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/market-movers/losers")
async def get_biggest_losers():
    """Get Biggest Stock Losers"""
    try:
        data = fetch_biggest_losers()
        if isinstance(data, dict) and 'error' in data:
            raise HTTPException(status_code=500, detail=data['error'])
        return {"losers": data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/market-movers/active")
async def get_most_actives():
    """Get Most Active Stocks"""
    try:
        data = fetch_most_actives()
        if isinstance(data, dict) and 'error' in data:
            raise HTTPException(status_code=500, detail=data['error'])
        return {"actives": data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/screener")
async def get_screener_stocks(
    marketCapMoreThan: Optional[int] = Query(None, description="Market Cap greater than"),
    sector: Optional[str] = Query(None, description="Sector filter"),
    limit: int = Query(20, description="Limit results")
):
    """Get Stocks via Screener"""
    try:
        data = fetch_stock_screener(marketCapMoreThan, sector, limit)
        if isinstance(data, dict) and 'error' in data:
            raise HTTPException(status_code=500, detail=data['error'])
        return {"screener": data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ============================================================================
# EARNINGS & NEWS ENDPOINTS
# ============================================================================

@app.get("/earnings-transcript-list/{symbol}")
async def get_earnings_transcripts_list(symbol: str):
    """Get list of available earnings transcripts"""
    symbol = symbol.upper()
    try:
        data = fetch_earnings_transcripts_list(symbol)
        if isinstance(data, dict) and 'error' in data:
            # Return empty list gracefully if it fails (restricted etc)
            print(f"Error fetching transcripts list for {symbol}: {data['error']}")
            return {"symbol": symbol, "transcripts": []}
        return {"symbol": symbol, "transcripts": data}
    except Exception as e:
        print(f"Exception in get_earnings_transcripts_list: {str(e)}")
        return {"symbol": symbol, "transcripts": []}

@app.get("/earnings-transcript/{symbol}")
async def get_earnings_transcript_content(
    symbol: str, 
    year: int = Query(..., description="Year"), 
    quarter: int = Query(..., description="Quarter")
):
    """Get content of specific transcript"""
    symbol = symbol.upper()
    try:
        data = fetch_earnings_transcript(symbol, year, quarter)
        if isinstance(data, dict) and 'error' in data:
            raise HTTPException(status_code=404, detail=data['error'])
        return {"symbol": symbol, "transcript": data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/news/fmp-articles")
async def get_fmp_articles_endpoint(
    page: int = Query(0, description="Page number"),
    limit: int = Query(20, description="Limit results")
):
    """Get FMP Articles"""
    try:
        data = fetch_fmp_articles(page, limit)
        if isinstance(data, dict) and 'error' in data:
             return []
        return data
    except Exception:
        return []

@app.get("/news/general")
async def get_general_news(page: int = 0, limit: int = 20):
    data = fetch_general_news(page, limit)
    return data

from modules.fmp import fetch_dividends

@app.get("/dividends/{symbol}")
async def get_dividends(symbol: str):
    symbol = symbol.upper()
    data = fetch_dividends(symbol)
    return data

@app.get("/stock-peers/{symbol}")
async def get_stock_peers_endpoint(symbol: str):
    symbol = symbol.upper()
    data = fetch_stock_peers(symbol)
    return data

# ============================================================================
# KNOWLEDGE BASE - RAG Value Investing Agent
# ============================================================================

from pydantic import BaseModel
from typing import List, Optional as OptionalType

class DocumentUpload(BaseModel):
    collection: str
    content: str
    title: OptionalType[str] = None
    symbol: OptionalType[str] = None
    tags: OptionalType[List[str]] = None

class SearchQuery(BaseModel):
    query: str
    collections: OptionalType[List[str]] = None
    n_results: int = 5

class ContextRequest(BaseModel):
    symbol: str
    company_name: str

@app.get("/knowledge/stats")
async def get_knowledge_stats():
    """Get statistics about the knowledge base"""
    try:
        from modules.knowledge_base import get_knowledge_base, reset_knowledge_base
        # Forzar reinicio para asegurar nueva configuración
        reset_knowledge_base()
        kb = get_knowledge_base()
        stats = kb.get_stats()
        return {"success": True, "stats": stats}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/knowledge/upload")
async def upload_document(doc: DocumentUpload):
    """Upload a document to the knowledge base"""
    try:
        from modules.knowledge_base import get_knowledge_base
        kb = get_knowledge_base()
        
        metadata = {}
        if doc.title:
            metadata["title"] = doc.title
        if doc.symbol:
            metadata["symbol"] = doc.symbol.upper()
        if doc.tags:
            metadata["tags"] = ",".join(doc.tags)
        
        # Colección unificada - ignoramos doc.collection
        result = kb.add_document(
            content=doc.content,
            metadata=metadata if metadata else None
        )
        return result
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/knowledge/search")
async def search_knowledge(query: SearchQuery):
    """Search the knowledge base"""
    try:
        from modules.knowledge_base import get_knowledge_base
        kb = get_knowledge_base()
        
        # Colección unificada - ignoramos query.collections
        results = kb.search(
            query=query.query,
            n_results=query.n_results
        )
        return {"success": True, "results": results, "count": len(results)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/knowledge/context")
async def get_analysis_context(req: ContextRequest):
    """Get relevant context for analyzing a specific stock"""
    try:
        from modules.knowledge_base import get_knowledge_base
        kb = get_knowledge_base()
        
        context = kb.get_context_for_analysis(
            symbol=req.symbol.upper(),
            company_name=req.company_name
        )
        return {"success": True, "context": context, "symbol": req.symbol.upper()}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/knowledge/list/{collection}")
async def list_documents(collection: str, limit: int = 50):
    """List documents in a collection"""
    try:
        from modules.knowledge_base import get_knowledge_base
        kb = get_knowledge_base()
        
        docs = kb.list_documents(limit=limit)
        return {"success": True, "documents": docs, "count": len(docs)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/knowledge/delete/{collection}/{document_id}")
async def delete_document(collection: str, document_id: str):
    """Delete a document from the knowledge base"""
    try:
        from modules.knowledge_base import get_knowledge_base
        kb = get_knowledge_base()
        
        success = kb.delete_document(collection, document_id)
        return {"success": success}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ============================================================================
# BATCH PDF UPLOAD
# ============================================================================

from fastapi import File, UploadFile
from typing import List as ListType

@app.post("/knowledge/upload-files")
async def upload_files(
    collection: str = "analyses",
    files: ListType[UploadFile] = File(...)
):
    """
    Upload multiple PDF/TXT files to the knowledge base
    Extracts text from PDFs and stores in ChromaDB
    """
    from modules.knowledge_base import get_knowledge_base
    import io
    
    kb = get_knowledge_base()
    results = []
    
    for file in files:
        try:
            filename = file.filename or "unknown"
            content_type = file.content_type or ""
            file_bytes = await file.read()
            
            # Extract text based on file type
            text_content = ""
            
            if filename.lower().endswith('.pdf') or 'pdf' in content_type.lower():
                # PDF extraction
                try:
                    from PyPDF2 import PdfReader
                    pdf_reader = PdfReader(io.BytesIO(file_bytes))
                    text_parts = []
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    text_content = "\n\n".join(text_parts)
                except Exception as pdf_error:
                    results.append({
                        "filename": filename,
                        "success": False,
                        "error": f"PDF parse error: {str(pdf_error)}"
                    })
                    continue
            
            elif filename.lower().endswith(('.xlsx', '.xls')):
                # Excel extraction
                try:
                    from openpyxl import load_workbook
                    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
                    text_parts = []
                    for sheet in wb.sheetnames:
                        ws = wb[sheet]
                        sheet_text = f"=== Hoja: {sheet} ===\n"
                        for row in ws.iter_rows(values_only=True):
                            row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                            if row_text.strip():
                                sheet_text += row_text + "\n"
                        text_parts.append(sheet_text)
                    text_content = "\n\n".join(text_parts)
                except Exception as excel_error:
                    results.append({
                        "filename": filename,
                        "success": False,
                        "error": f"Excel parse error: {str(excel_error)}"
                    })
                    continue
            
            elif filename.lower().endswith('.docx'):
                # Word extraction
                try:
                    from docx import Document
                    doc = Document(io.BytesIO(file_bytes))
                    text_parts = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            text_parts.append(para.text)
                    # También extraer tablas
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = " | ".join([cell.text for cell in row.cells])
                            if row_text.strip():
                                text_parts.append(row_text)
                    text_content = "\n\n".join(text_parts)
                except Exception as docx_error:
                    results.append({
                        "filename": filename,
                        "success": False,
                        "error": f"Word parse error: {str(docx_error)}"
                    })
                    continue
                    
            elif filename.lower().endswith(('.txt', '.md')):
                # Plain text files
                try:
                    text_content = file_bytes.decode('utf-8')
                except UnicodeDecodeError:
                    text_content = file_bytes.decode('latin-1')
            else:
                results.append({
                    "filename": filename,
                    "success": False,
                    "error": "Tipo no soportado. Usa PDF, Excel, Word, TXT o MD."
                })
                continue
            
            # Skip empty files
            if not text_content.strip():
                results.append({
                    "filename": filename,
                    "success": False,
                    "error": "File is empty or could not extract text"
                })
                continue
            
            # Add to knowledge base (colección unificada)
            result = kb.add_document(
                content=text_content,
                metadata={"title": filename, "source": "file_upload"}
            )
            
            results.append({
                "filename": filename,
                "success": True,
                "document_id": result.get("document_id"),
                "chunks_added": result.get("chunks_added", 0)
            })
            
        except Exception as e:
            results.append({
                "filename": file.filename or "unknown",
                "success": False,
                "error": str(e)
            })
    
    successful = sum(1 for r in results if r.get("success"))
    return {
        "success": True,
        "total_files": len(files),
        "successful": successful,
        "failed": len(files) - successful,
        "results": results
    }
