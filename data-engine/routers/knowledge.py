"""Knowledge base API routes."""

from __future__ import annotations

from typing import List, Optional

from fastapi import APIRouter, File, HTTPException, UploadFile
from pydantic import BaseModel

router = APIRouter(prefix="/knowledge", tags=["knowledge"])


class DocumentUpload(BaseModel):
    collection: str
    content: str
    title: Optional[str] = None
    symbol: Optional[str] = None
    tags: Optional[List[str]] = None


class SearchQuery(BaseModel):
    query: str
    collections: Optional[List[str]] = None
    n_results: int = 5
    symbol: Optional[str] = None


class ContextRequest(BaseModel):
    symbol: str
    company_name: str


@router.get("/stats")
async def get_knowledge_stats():
    try:
        from modules.knowledge_base import get_knowledge_base, reset_knowledge_base

        reset_knowledge_base()
        kb = get_knowledge_base()
        return {"success": True, "stats": kb.get_stats()}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/upload")
async def upload_document(doc: DocumentUpload):
    try:
        from modules.knowledge_base import get_knowledge_base

        metadata = {}
        if doc.title:
            metadata["title"] = doc.title
        if doc.symbol:
            metadata["symbol"] = doc.symbol.upper()
        if doc.tags:
            metadata["tags"] = ",".join(doc.tags)
        result = get_knowledge_base().add_document(
            content=doc.content,
            metadata=metadata if metadata else None,
        )
        return result
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/search")
async def search_knowledge(query: SearchQuery):
    try:
        from modules.knowledge_base import get_knowledge_base

        results = get_knowledge_base().search(
            query=query.query,
            n_results=query.n_results,
            symbol=query.symbol,
        )
        return {"success": True, "results": results, "count": len(results)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/context")
async def get_analysis_context(req: ContextRequest):
    try:
        from modules.knowledge_base import get_knowledge_base

        context = get_knowledge_base().get_context_for_analysis(
            symbol=req.symbol.upper(),
            company_name=req.company_name,
        )
        return {"success": True, "context": context, "symbol": req.symbol.upper()}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/list/{collection}")
async def list_documents(collection: str, limit: int = 50):
    try:
        from modules.knowledge_base import get_knowledge_base

        docs = get_knowledge_base().list_documents(limit=limit)
        return {"success": True, "documents": docs, "count": len(docs)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/delete/{collection}/{document_id}")
async def delete_document(collection: str, document_id: str):
    try:
        from modules.knowledge_base import get_knowledge_base

        success = get_knowledge_base().delete_document(document_id)
        return {"success": success}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/upload-files")
async def upload_files(collection: str = "analyses", files: List[UploadFile] = File(...)):
    from modules.knowledge_base import get_knowledge_base
    import io

    kb = get_knowledge_base()
    results = []
    for file in files:
        try:
            filename = file.filename or "unknown"
            file_bytes = await file.read()
            text_content = ""
            if filename.lower().endswith(".pdf"):
                from PyPDF2 import PdfReader

                reader = PdfReader(io.BytesIO(file_bytes))
                text_content = "\n\n".join([page.extract_text() or "" for page in reader.pages])
            elif filename.lower().endswith((".xlsx", ".xls")):
                from openpyxl import load_workbook

                workbook = load_workbook(io.BytesIO(file_bytes), data_only=True)
                text_parts = []
                for sheet in workbook.sheetnames:
                    worksheet = workbook[sheet]
                    rows = []
                    for row in worksheet.iter_rows(values_only=True):
                        row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                        if row_text.strip():
                            rows.append(row_text)
                    if rows:
                        text_parts.append(f"=== Hoja: {sheet} ===\n" + "\n".join(rows))
                text_content = "\n\n".join(text_parts)
            elif filename.lower().endswith(".docx"):
                from docx import Document

                document = Document(io.BytesIO(file_bytes))
                text_parts = [para.text for para in document.paragraphs if para.text.strip()]
                for table in document.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text for cell in row.cells])
                        if row_text.strip():
                            text_parts.append(row_text)
                text_content = "\n\n".join(text_parts)
            elif filename.lower().endswith((".txt", ".md")):
                try:
                    text_content = file_bytes.decode("utf-8")
                except UnicodeDecodeError:
                    text_content = file_bytes.decode("latin-1")
            else:
                results.append({"filename": filename, "success": False, "error": "Tipo no soportado. Usa PDF, Excel, Word, TXT o MD."})
                continue
            if not text_content.strip():
                results.append({"filename": filename, "success": False, "error": "File is empty"})
                continue
            result = kb.add_document(content=text_content, metadata={"title": filename, "source": "file_upload"})
            results.append({
                "filename": filename,
                "success": True,
                "document_id": result.get("document_id"),
                "chunks_added": result.get("chunks_added", 0),
            })
        except Exception as e:
            results.append({"filename": file.filename or "unknown", "success": False, "error": str(e)})
    return {"success": True, "results": results}
